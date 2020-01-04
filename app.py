# from flask import Flask, render_template, request, send_from_directory
# from werkzeug.utils import secure_filename
# import os
# import shutil

# import re
# from pdfminer.pdfparser import PDFParser, PDFDocument
# from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
# from pdfminer.converter import PDFPageAggregator
# from pdfminer.layout import LAParams, LTTextBox, LTTextLine	

# import docx
# import os
# import subprocess
# import sys


# from attention import AttentionLayer

# import re
# import random
# import numpy as np
# import pandas as pd 
# import tensorflow as tf
# from bs4 import BeautifulSoup 
# from keras.preprocessing.text import Tokenizer 
# from keras.preprocessing.sequence import pad_sequences
# from nltk.corpus import stopwords
# from tensorflow.keras.layers import Input, LSTM, Embedding, Dense, Concatenate, TimeDistributed, Bidirectional
# from tensorflow.keras.models import Model
# from tensorflow.keras.callbacks import EarlyStopping
# import warnings
# pd.set_option("display.max_colwidth", 200)
# warnings.filterwarnings("ignore")
# import nltk
# from nltk.tokenize import sent_tokenize
# from sklearn.metrics.pairwise import cosine_similarity
# import networkx as nx
# from tensorflow.keras.models import load_model
# import pickle



# import re
# import nltk
# from nltk.tokenize import word_tokenize
# from nltk.corpus import stopwords 
# import random


# # print('This is error output', file=sys.stderr)
# # print('This is standard output', file=sys.stdout)


# app = Flask(__name__)
# UPLOAD_FOLDER = '/Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/pdf/'
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


# global word_embeddings
# word_embeddings = {}	

# f = open('glove.6B/glove.6B.100d.txt', encoding='utf-8')
# for line in f:
#     values = line.split()
#     word = values[0]
#     coefs = np.asarray(values[1:], dtype='float32')
#     word_embeddings[word] = coefs
# f.close()


# @app.route("/", methods=["GET","POST"])
# def mainpage():
# 	if request.method == 'POST':
# 		try:
# 			abstract = request.form.get("Summary")
# 			#print(abstract)
			
# 			if(abstract != None):
# 				abstract = (str(abstract).lstrip()).rstrip()
# 				match_key = re.search("(Keywords|KEYWORDS)((.|\n)*)",abstract)				

# 				if(match_key == None):
# 				    keywords = ""
# 				else:
# 				    keywords = (str(match_key.group(0)).lstrip()).rstrip()

# 			else:
# 				file = request.files['Files']
# 				print(file)	

# 				filename = secure_filename(file.filename)
# 				file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))	

				

# 				########## Code for Extracting Abstract and Keywords from PDF   ###########				

# 				fp = open(os.path.join(app.config['UPLOAD_FOLDER'], filename), 'rb')				

# 				f1 = open('Test_doc.txt', 'w+')
# 				parser = PDFParser(fp)
# 				doc = PDFDocument()
# 				parser.set_document(doc)
# 				doc.set_parser(parser)
# 				doc.initialize('')
# 				rsrcmgr = PDFResourceManager()
# 				laparams = LAParams()
# 				device = PDFPageAggregator(rsrcmgr, laparams=laparams)
# 				interpreter = PDFPageInterpreter(rsrcmgr, device)				

# 				extracted_text = ""
# 				for page in doc.get_pages():
# 				    interpreter.process_page(page)
# 				    layout = device.get_result()
# 				    for lt_obj in layout:
# 				        if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
# 				            extracted_text += lt_obj.get_text()
# 				            f1.write(lt_obj.get_text())
# 				            #print(lt_obj.get_text())				

# 				f1.close()
# 				fp.close()				

# 				f = open('Test_doc.txt',"r")
# 				f1 = open("Abstract.txt","w+")
# 				f2 = open("Keywords.txt","w+")				

# 				string = f.read()
# 				match = re.search("(Abstract|ABSTRACT)((.|\n)*)(I|1)\s*\.*?\s*(Introduction|INTRODUCTION|IN TRODUC T ION)",string)
# 				abstract = (str(match.group(2)).lstrip()).rstrip()
# 				match_key = re.search("(Keywords|KEYWORDS)((.|\n)*)",abstract)				

# 				if(match_key == None):
# 				    keywords = ""
# 				else:
# 				    keywords = (str(match_key.group(0)).lstrip()).rstrip()				

# 				f2.write(keywords)
# 				f1.write(abstract)
# 				f2.close()
# 				f1.close()
# 				f.close()				

# 				################   ------------    #################				

# 				################    Code for Extracting title from the PDF file     ##################						

# 				for top, dirs, files in os.walk('/Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/pdf/'):
# 				    for filename in files:
# 				        if filename.endswith('.pdf'):
# 				            abspath = os.path.join(top, filename)
# 				            subprocess.call('./soffice --convert-to html "{}" --outdir /Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/html'.format(abspath), shell=True)				

# 				for top, dirs, files in os.walk('/Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/html/'):
# 				    for filename in files:
# 				        if filename.endswith('.html'):
# 				            abspath = os.path.join(top, filename)
# 				            subprocess.call('./soffice --convert-to docx:"MS Word 2007 XML" "{}" --outdir /Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/docx/'.format(abspath), shell=True)				

# 				for top, dirs, files in os.walk('/Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/docx/'):
# 				    for filename in files:
# 				        if filename.endswith('.docx'):
# 				            abspath = os.path.join(top, filename)
# 				            document = docx.Document(abspath)
# 				            bolds=[]
# 				            italics=[]
# 				            count = 0
# 				            count_real = 0
# 				            temp = ""
# 				            flag = True
# 				            for para in document.paragraphs:
# 				            	if(not flag):
# 				            		break;
# 				            	for run in para.runs:
# 				            		if(run.text!=""):
# 				            			count_real+=1
# 				            			temp+=run.text + " "
# 				            		if(count_real==2):
# 				            			temp1 = temp
# 				            		if(not flag):
# 				            			break;
# 				            		if(count==1):
# 				            			flag = False
# 				            		if run.bold:
# 				            			if(run.text!=""):
# 				            				bolds.append(run.text)
# 				            				count+=1				

# 				            boltalic_Dict={'bold_phrases':bolds}
# 				            title = ""		
# 				            for i in bolds:
# 				            	if(i!=""):
# 				            		title += i+" "
# 				            if(len(title)<2):
# 				            	title=temp1	

# 				print(title)
# 				f3 = open("Ground_Title.txt","w+")
# 				f3.write(title)
# 				f3.close()


# 			# ##############      -----------------       #############

# 			##############		Code for Generating Title from Abstract 	###########

# 			contraction_mapping = {"ain't": "is not", "aren't": "are not","can't": "cannot", "'cause": "because", "could've": "could have", "couldn't": "could not",						

# 						                           "didn't": "did not", "doesn't": "does not", "don't": "do not", "hadn't": "had not", "hasn't": "has not", "haven't": "have not",						

# 						                           "he'd": "he would","he'll": "he will", "he's": "he is", "how'd": "how did", "how'd'y": "how do you", "how'll": "how will", "how's": "how is",						

# 						                           "I'd": "I would", "I'd've": "I would have", "I'll": "I will", "I'll've": "I will have","I'm": "I am", "I've": "I have", "i'd": "i would",						

# 						                           "i'd've": "i would have", "i'll": "i will",  "i'll've": "i will have","i'm": "i am", "i've": "i have", "isn't": "is not", "it'd": "it would",						

# 						                           "it'd've": "it would have", "it'll": "it will", "it'll've": "it will have","it's": "it is", "let's": "let us", "ma'am": "madam",						

# 						                           "mayn't": "may not", "might've": "might have","mightn't": "might not","mightn't've": "might not have", "must've": "must have",						

# 						                           "mustn't": "must not", "mustn't've": "must not have", "needn't": "need not", "needn't've": "need not have","o'clock": "of the clock",						

# 						                           "oughtn't": "ought not", "oughtn't've": "ought not have", "shan't": "shall not", "sha'n't": "shall not", "shan't've": "shall not have",						

# 						                           "she'd": "she would", "she'd've": "she would have", "she'll": "she will", "she'll've": "she will have", "she's": "she is",						

# 						                           "should've": "should have", "shouldn't": "should not", "shouldn't've": "should not have", "so've": "so have","so's": "so as",						

# 						                           "this's": "this is","that'd": "that would", "that'd've": "that would have", "that's": "that is", "there'd": "there would",						

# 						                           "there'd've": "there would have", "there's": "there is", "here's": "here is","they'd": "they would", "they'd've": "they would have",						

# 						                           "they'll": "they will", "they'll've": "they will have", "they're": "they are", "they've": "they have", "to've": "to have",						

# 						                           "wasn't": "was not", "we'd": "we would", "we'd've": "we would have", "we'll": "we will", "we'll've": "we will have", "we're": "we are",						

# 						                           "we've": "we have", "weren't": "were not", "what'll": "what will", "what'll've": "what will have", "what're": "what are",						

# 						                           "what's": "what is", "what've": "what have", "when's": "when is", "when've": "when have", "where'd": "where did", "where's": "where is",						

# 						                           "where've": "where have", "who'll": "who will", "who'll've": "who will have", "who's": "who is", "who've": "who have",						

# 						                           "why's": "why is", "why've": "why have", "will've": "will have", "won't": "will not", "won't've": "will not have",						

# 						                           "would've": "would have", "wouldn't": "would not", "wouldn't've": "would not have", "y'all": "you all",						

# 						                           "y'all'd": "you all would","y'all'd've": "you all would have","y'all're": "you all are","y'all've": "you all have",						

# 						                           "you'd": "you would", "you'd've": "you would have", "you'll": "you will", "you'll've": "you will have",						

# 						                           "you're": "you are", "you've": "you have"}						

# 			stop_words = set(stopwords.words('english'))
			
# 			# word_embeddings = {}			

# 			# f = open('glove.6B/glove.6B.100d.txt', encoding='utf-8')
# 			# for line in f:
# 			#     values = line.split()
# 			#     word = values[0]
# 			#     coefs = np.asarray(values[1:], dtype='float32')
# 			#     word_embeddings[word] = coefs
# 			# f.close()						

# 			encoderModel = load_model("titlegen.h5", custom_objects={'AttentionLayer': AttentionLayer})
# 			decoderModel = load_model("titlegenPredict.h5", custom_objects={'AttentionLayer': AttentionLayer})						

# 			with open('xtokenizer.pickle', 'rb') as handle:
# 			    x_tokenizer = pickle.load(handle)
# 			with open('ytokenizer.pickle', 'rb') as handle:
# 			    y_tokenizer = pickle.load(handle)						

# 			reverse_target_word_index=y_tokenizer.index_word
# 			reverse_source_word_index=x_tokenizer.index_word
# 			target_word_index=y_tokenizer.word_index						
			

# 			max_text_len=200
# 			max_summary_len=15
# 			no_of_extracted_sentences = 10

# 			def text_cleaner(text):
# 			    newString = text.lower()
# 			    newString = BeautifulSoup(newString, "lxml").text
# 			    newString = re.sub(r'\([^)]*\)', '', newString)
# 			    newString = re.sub('"','', newString)
# 			    newString = ' '.join([contraction_mapping[t] if t in contraction_mapping else t for t in newString.split(" ")])    
# 			    newString = re.sub(r"'s\b","",newString)
# 			    newString = re.sub("[^a-zA-Z]", " ", newString) 
# 			    tokens = [w for w in newString.split() if not w in stop_words]
# 			    long_words=[]
# 			    for i in tokens:
# 			        if len(i)>=3:                  #removing short word
# 			            long_words.append(i)   
# 			    return (" ".join(long_words)).strip()

# 			def extractText(text):
# 			    sentences = sent_tokenize(text)
# 			    clean_sentences = list()			

# 			    for sentence in sentences:
# 			        clean_sentences.append(text_cleaner(sentence))			

# 			    sentence_vectors = []
# 			    for i in clean_sentences:
# 			        if len(i) != 0:
# 			            v = sum([word_embeddings.get(w, np.zeros((100,))) for w in i.split()])/(len(i.split())+0.001)
# 			        else:
# 			            v = np.zeros((100,))
# 			        sentence_vectors.append(v)			

# 			    # similarity matrix
# 			    sim_mat = np.zeros([len(sentences), len(sentences)])			

# 			    for i in range(len(sentences)):
# 			        for j in range(len(sentences)):
# 			            if i != j:
# 			                sim_mat[i][j] = cosine_similarity(sentence_vectors[i].reshape(1,100), sentence_vectors[j].reshape(1,100))[0,0]			

# 			    nx_graph = nx.from_numpy_array(sim_mat)			

# 			    try:
# 			        scores = nx.pagerank(nx_graph)
# 			    except:
# 			        exit(1)			

# 			    ranked_sentences = sorted(((scores[i],s) for i,s in enumerate(sentences)), reverse=True)			

# 			    es = list()
# 			    for i in range(min(no_of_extracted_sentences, len(ranked_sentences))):
# 			        es.append(ranked_sentences[i][1])			

# 			    extracted_text = " ".join(es)			

# 			    return extracted_text

# 			def decode_sequence(input_seq):
# 			    # Encode the input as state vectors.
# 			    e_out, e_h, e_c = encoderModel.predict(input_seq)
			    
# 			    # Generate empty target sequence of length 1.
# 			    target_seq = np.zeros((1,1))
			    
# 			    # Populate the first word of target sequence with the start word.
# 			    target_seq[0, 0] = target_word_index['sostok']			

# 			    stop_condition = False
# 			    decoded_sentence = ''
# 			    while not stop_condition:
			      
# 			        output_tokens, h, c = decoderModel.predict([target_seq] + [e_out, e_h, e_c])			

# 			        # Sample a token
# 			        sampled_token_index = np.argmax(output_tokens[0, -1, :])
# 			        sampled_token = reverse_target_word_index[sampled_token_index]
			        
# 			        if(sampled_token!='eostok'):
# 			            decoded_sentence += ' '+sampled_token			

# 			        # Exit condition: either hit max length or find stop word.
# 			        if (sampled_token == 'eostok'  or len(decoded_sentence.split()) >= (max_summary_len-1)):
# 			            stop_condition = True			

# 			        # Update the target sequence (of length 1).
# 			        target_seq = np.zeros((1,1))
# 			        target_seq[0, 0] = sampled_token_index			

# 			        # Update internal states
# 			        e_h, e_c = h, c			

# 			    return decoded_sentence

# 			def seq2summary(input_seq):
# 			    newString=''
# 			    for i in input_seq:
# 			        if((i!=0 and i!=target_word_index['sostok']) and i!=target_word_index['eostok']):
# 			            newString=newString+reverse_target_word_index[i]+' '
# 			    return newString

# 			def seq2text(input_seq):
# 			    newString=''
# 			    for i in input_seq:
# 			        if(i!=0):
# 			            newString=newString+reverse_source_word_index[i]+' '
# 			    return newString


# 			extracted_text = extractText(abstract)
# 			cleaned_text = text_cleaner(extracted_text)			
# 			x_val_seq = x_tokenizer.texts_to_sequences([cleaned_text])
# 			text_encodings = pad_sequences(x_val_seq, maxlen=max_text_len, padding='post')[0]			
# 			resultTitle = decode_sequence(text_encodings.reshape(1,max_text_len))
# 			print(resultTitle)

# 			############	---------------		###############

# 			############	Code for Making the generated title Catchy 	###############

# 			def count(string):
# 				return sum(1 for c in string if c.isupper())			

# 			def catchy1(t,text,key):
# 				s = t.split()
# 				r = []
# 				flag = 0
# 				r.append(s[0])
# 				for i in range (1,len(s)):
# 					if s[i] == s[i-1]:
# 						continue
# 					else:
# 						r.append(s[i])
# 				for i in range (len(r)):
# 					r[i] = r[i][0].upper()+r[i][1:]
# 				stop_words = set(stopwords.words('english')) 
# 				text = re.sub(r'[^A-Za-z\s]',"",text)
# 				text = re.sub(r'https\S+', "", text)
# 				text = text.lstrip()
# 				text = text.strip()
# 				list1 = word_tokenize(text)
# 				fil_sent = [w for w in list1 if not w in stop_words]			

# 				final = []
# 				c = {}
# 				for i in fil_sent:
# 					if i[1:] != i[1:].lower() and count(i)>3:
# 						final.append(i)
# 						if i not in c:
# 							c[i]=1
# 						else:
# 							c[i] += 1
# 				ans = ""
# 				if len(c)!=0:
# 					max_value = max(c.values())
# 					last = [k for k,v in c.items() if v == max_value]
# 					length = []
# 					if len(last)>=1:
# 						for i in range (len(last)):
# 							if last[i] == last[i].upper():
# 								flag = 1
# 								ans = last[i]
# 								length.append(len(last[i]))
# 							else:
# 								length.append(len(last[i]))
# 						if ans == "":
# 							flag = 1
# 							ans = last[length.index(max(length))]
# 				final = []
# 				if ans == "":
# 					if len(key)>0:
# 						flag = 2
# 						pre = re.split('-|;|,|—|:',key)
# 						for i in range (len(pre)):
# 							if pre[i] != "" and pre[i] !="keywords" and pre[i]!="Keywords" and pre[i]!="KEYWORDS":
# 								final.append(pre[i])
# 						for i in range(len(final)):
# 							final[i] = final[i].lstrip()
# 							final[i] = final[i].rstrip()
# 						print(final)
# 						z = random.randrange(0,len(final))
# 						ans = final[z]
# 						ans = ans.replace(" ","")			

# 				ans1 = ""
# 				for i in range (len(r)):
# 					if i == 0:
# 						ans1 = ans1+r[i]
# 					else:
# 						ans1 = ans1+" "+r[i]
# 				if flag == 2:
# 					ans1 = "#"+ans+": "+ans1
# 				elif flag == 0:
# 					ans1 = ans1
# 				else:
# 					ans1 = ans+": "+ans1
# 				return (ans1)			

# 			catchy_title = catchy1(resultTitle,abstract,keywords)
# 			print(catchy_title)

# 			f3 = open("Predicted_Title.txt","w+")
# 			f3.write(catchy_title)
# 			f3.close()

# 			#############	--------------	#############

# 			#############	Code for Generating the metric for Catchiness	############

# 			with open('xtokenizer.pickle', 'rb') as handle:
# 			    x_tokenizer = pickle.load(handle)
# 			with open('ytokenizer.pickle', 'rb') as handle:
# 			    y_tokenizer = pickle.load(handle)						

# 			dict1 = y_tokenizer.word_counts
# 			dict2 = x_tokenizer.word_counts						

# 			def merge_two_dicts(x, y):
# 				z = x.copy()   
# 				z.update(y)    
# 				return z						

# 			dict_new = merge_two_dicts(dict1,dict2)			

# 			def catchyscore(ground,catchy):
# 				stopwords = nltk.corpus.stopwords.words('english')				

# 				ground_new = []
# 				catchy_new = []				

# 				ground = ground.split(" ")
# 				catchy = catchy.split(" ")			

# 				for i in ground:
# 					if(i not in stopwords):
# 						ground_new.append(i)			

# 				for i in catchy:
# 					if i not in stopwords:
# 						catchy_new.append(i)				

# 				ground_score = 0	# default values
# 				catchy_score = random.choice([-1,0,1,2,3])	# default values			

# 				for i in range(len(ground_new)):
# 					if ground_new[i] in dict_new:
# 						ground_score+=dict_new[i]				

# 				for i in range(len(catchy_new)):
# 					if catchy_new[i] in dict_new:
# 						catchy_score+=dict_new[i]				

# 				final = ground_score - catchy_score
# 				return final		

# 			if(abstract!=None):
# 				catch_score = catchyscore(title,catchy_title)
# 				print(catch_score)

# 			#############	--------------	###############

# 		except:
# 			return render_template("mainpage.html")

# 	return render_template("mainpage.html")

# @app.route("/index", methods=["GET","POST"])
# def index():
#  	return "Hello, world!"

# # @app.route("/hello", methods=["POST"])
# # def hello():
# # 	return "Hello"

# # @app.route("/graph")
# # def graph():

# # @app.route("/title_show", methods=["GET","POST"])
# # def title_show():



# app.config['ENV'] = 'development'
# app.config['DEBUG'] = True
# app.config['TESTING'] = True
# if __name__ == '__main__':
# 	app.run(debug=False)






from flask import Flask, render_template, request, send_from_directory, redirect
from werkzeug.utils import secure_filename
import os
import shutil

import re
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine	

import docx
import os
import subprocess
import sys


from attention import AttentionLayer

import re
import random
import numpy as np
import pandas as pd 
import tensorflow as tf
from bs4 import BeautifulSoup 
from keras.preprocessing.text import Tokenizer 
from keras.preprocessing.sequence import pad_sequences
from nltk.corpus import stopwords
from tensorflow.keras.layers import Input, LSTM, Embedding, Dense, Concatenate, TimeDistributed, Bidirectional
from tensorflow.keras.models import Model
from tensorflow.keras.callbacks import EarlyStopping
import warnings
pd.set_option("display.max_colwidth", 200)
warnings.filterwarnings("ignore")
import nltk
from nltk.tokenize import sent_tokenize
from sklearn.metrics.pairwise import cosine_similarity
import networkx as nx
from tensorflow.keras.models import load_model
import pickle



import re
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords 
import random


# print('This is error output', file=sys.stderr)
# print('This is standard output', file=sys.stdout)


app = Flask(__name__)
UPLOAD_FOLDER = '/Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/pdf/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


global word_embeddings
word_embeddings = {}	

f = open('glove.6B/glove.6B.100d.txt', encoding='utf-8')
for line in f:
    values = line.split()
    word = values[0]
    coefs = np.asarray(values[1:], dtype='float32')
    word_embeddings[word] = coefs
f.close()

contraction_mapping = {"ain't": "is not", "aren't": "are not","can't": "cannot", "'cause": "because", "could've": "could have", "couldn't": "could not",						

						                           "didn't": "did not", "doesn't": "does not", "don't": "do not", "hadn't": "had not", "hasn't": "has not", "haven't": "have not",						

						                           "he'd": "he would","he'll": "he will", "he's": "he is", "how'd": "how did", "how'd'y": "how do you", "how'll": "how will", "how's": "how is",						

						                           "I'd": "I would", "I'd've": "I would have", "I'll": "I will", "I'll've": "I will have","I'm": "I am", "I've": "I have", "i'd": "i would",						

						                           "i'd've": "i would have", "i'll": "i will",  "i'll've": "i will have","i'm": "i am", "i've": "i have", "isn't": "is not", "it'd": "it would",						

						                           "it'd've": "it would have", "it'll": "it will", "it'll've": "it will have","it's": "it is", "let's": "let us", "ma'am": "madam",						

						                           "mayn't": "may not", "might've": "might have","mightn't": "might not","mightn't've": "might not have", "must've": "must have",						

						                           "mustn't": "must not", "mustn't've": "must not have", "needn't": "need not", "needn't've": "need not have","o'clock": "of the clock",						

						                           "oughtn't": "ought not", "oughtn't've": "ought not have", "shan't": "shall not", "sha'n't": "shall not", "shan't've": "shall not have",						

						                           "she'd": "she would", "she'd've": "she would have", "she'll": "she will", "she'll've": "she will have", "she's": "she is",						

						                           "should've": "should have", "shouldn't": "should not", "shouldn't've": "should not have", "so've": "so have","so's": "so as",						

						                           "this's": "this is","that'd": "that would", "that'd've": "that would have", "that's": "that is", "there'd": "there would",						

						                           "there'd've": "there would have", "there's": "there is", "here's": "here is","they'd": "they would", "they'd've": "they would have",						

						                           "they'll": "they will", "they'll've": "they will have", "they're": "they are", "they've": "they have", "to've": "to have",						

						                           "wasn't": "was not", "we'd": "we would", "we'd've": "we would have", "we'll": "we will", "we'll've": "we will have", "we're": "we are",						

						                           "we've": "we have", "weren't": "were not", "what'll": "what will", "what'll've": "what will have", "what're": "what are",						

						                           "what's": "what is", "what've": "what have", "when's": "when is", "when've": "when have", "where'd": "where did", "where's": "where is",						

						                           "where've": "where have", "who'll": "who will", "who'll've": "who will have", "who's": "who is", "who've": "who have",						

						                           "why's": "why is", "why've": "why have", "will've": "will have", "won't": "will not", "won't've": "will not have",						

						                           "would've": "would have", "wouldn't": "would not", "wouldn't've": "would not have", "y'all": "you all",						

						                           "y'all'd": "you all would","y'all'd've": "you all would have","y'all're": "you all are","y'all've": "you all have",						

						                           "you'd": "you would", "you'd've": "you would have", "you'll": "you will", "you'll've": "you will have",						

						                           "you're": "you are", "you've": "you have"}						

stop_words = set(stopwords.words('english'))

# @app.route("/")
# def mainpage():
# 	return render_template("mainpage.html")


@app.route("/", methods=["GET","POST"])
def mainpage():
	orgtit = 0
	if request.method == 'POST':
		try:
			abstract = request.form.get("Summary")
			#print(abstract)
			title = None
			if(abstract != None):
				abstract = (str(abstract).lstrip()).rstrip()
				match_key = re.search("(Keywords|KEYWORDS)((.|\n)*)",abstract)				

				if(match_key == None):
				    keywords = ""
				else:
				    keywords = (str(match_key.group(0)).lstrip()).rstrip()

			else:
				file = request.files['Files']
				print(file)	

				filename = secure_filename(file.filename)
				file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))	

				

				########## Code for Extracting Abstract and Keywords from PDF   ###########				

				fp = open(os.path.join(app.config['UPLOAD_FOLDER'], filename), 'rb')				

				f1 = open('Test_doc.txt', 'w+')
				parser = PDFParser(fp)
				doc = PDFDocument()
				parser.set_document(doc)
				doc.set_parser(parser)
				doc.initialize('')
				rsrcmgr = PDFResourceManager()
				laparams = LAParams()
				device = PDFPageAggregator(rsrcmgr, laparams=laparams)
				interpreter = PDFPageInterpreter(rsrcmgr, device)				

				extracted_text = ""
				for page in doc.get_pages():
				    interpreter.process_page(page)
				    layout = device.get_result()
				    for lt_obj in layout:
				        if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
				            extracted_text += lt_obj.get_text()
				            f1.write(lt_obj.get_text())
				            #print(lt_obj.get_text())				

				f1.close()
				fp.close()				

				f = open('Test_doc.txt',"r")
				f1 = open("Abstract.txt","w+")
				f2 = open("Keywords.txt","w+")				

				string = f.read()
				match = re.search("(Abstract|ABSTRACT)((.|\n)*)(I|1)\s*\.*?\s*(Introduction|INTRODUCTION|IN TRODUC T ION)",string)
				abstract = (str(match.group(2)).lstrip()).rstrip()
				match_key = re.search("(Keywords|KEYWORDS)((.|\n)*)",abstract)				

				if(match_key == None):
				    keywords = ""
				else:
				    keywords = (str(match_key.group(0)).lstrip()).rstrip()				

				f2.write(keywords)
				f1.write(abstract)
				f2.close()
				f1.close()
				f.close()				
			
				################    Code for Extracting title from the PDF file     ##################						

				for top, dirs, files in os.walk('/Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/pdf/'):
				    for filename in files:
				        if filename.endswith('.pdf'):
				            abspath = os.path.join(top, filename)
				            subprocess.call('./soffice --convert-to html "{}" --outdir /Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/html'.format(abspath), shell=True)				

				for top, dirs, files in os.walk('/Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/html/'):
				    for filename in files:
				        if filename.endswith('.html'):
				            abspath = os.path.join(top, filename)
				            subprocess.call('./soffice --convert-to docx:"MS Word 2007 XML" "{}" --outdir /Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/docx/'.format(abspath), shell=True)				

				for top, dirs, files in os.walk('/Users/anubhavjain/Desktop/NLP_Pdf_extract/paper1/docx/'):
				    for filename in files:
				        if filename.endswith('.docx'):
				            abspath = os.path.join(top, filename)
				            document = docx.Document(abspath)
				            bolds=[]
				            italics=[]
				            count = 0
				            count_real = 0
				            temp = ""
				            flag = True
				            for para in document.paragraphs:
				            	if(not flag):
				            		break;
				            	for run in para.runs:
				            		if(run.text!=""):
				            			count_real+=1
				            			temp+=run.text + " "
				            		if(count_real==2):
				            			temp1 = temp
				            		if(not flag):
				            			break;
				            		if(count==1):
				            			flag = False
				            		if run.bold:
				            			if(run.text!=""):
				            				bolds.append(run.text)
				            				count+=1				

				            boltalic_Dict={'bold_phrases':bolds}
				            title = ""		
				            for i in bolds:
				            	if(i!=""):
				            		title += i+" "
				            if(len(title)<2):
				            	title=temp1	

				print(title)
				f3 = open("Ground_Title.txt","w+")
				f3.write(title)
				f3.close()
				orgtit = 1


			# ##############      -----------------       #############

			encoderModel = load_model("titlegen.h5", custom_objects={'AttentionLayer': AttentionLayer})
			decoderModel = load_model("titlegenPredict.h5", custom_objects={'AttentionLayer': AttentionLayer})						

			with open('xtokenizer.pickle', 'rb') as handle:
			    x_tokenizer = pickle.load(handle)
			with open('ytokenizer.pickle', 'rb') as handle:
			    y_tokenizer = pickle.load(handle)						

			reverse_target_word_index=y_tokenizer.index_word
			reverse_source_word_index=x_tokenizer.index_word
			target_word_index=y_tokenizer.word_index						
			

			max_text_len=200
			max_summary_len=15
			no_of_extracted_sentences = 10

			def text_cleaner(text):
			    newString = text.lower()
			    newString = BeautifulSoup(newString, "lxml").text
			    newString = re.sub(r'\([^)]*\)', '', newString)
			    newString = re.sub('"','', newString)
			    newString = ' '.join([contraction_mapping[t] if t in contraction_mapping else t for t in newString.split(" ")])    
			    newString = re.sub(r"'s\b","",newString)
			    newString = re.sub("[^a-zA-Z]", " ", newString) 
			    tokens = [w for w in newString.split() if not w in stop_words]
			    long_words=[]
			    for i in tokens:
			        if len(i)>=3:                  #removing short word
			            long_words.append(i)   
			    return (" ".join(long_words)).strip()

			def extractText(text):
			    sentences = sent_tokenize(text)
			    clean_sentences = list()			

			    for sentence in sentences:
			        clean_sentences.append(text_cleaner(sentence))			

			    sentence_vectors = []
			    for i in clean_sentences:
			        if len(i) != 0:
			            v = sum([word_embeddings.get(w, np.zeros((100,))) for w in i.split()])/(len(i.split())+0.001)
			        else:
			            v = np.zeros((100,))
			        sentence_vectors.append(v)			

			    # similarity matrix
			    sim_mat = np.zeros([len(sentences), len(sentences)])			

			    for i in range(len(sentences)):
			        for j in range(len(sentences)):
			            if i != j:
			                sim_mat[i][j] = cosine_similarity(sentence_vectors[i].reshape(1,100), sentence_vectors[j].reshape(1,100))[0,0]			

			    nx_graph = nx.from_numpy_array(sim_mat)			

			    try:
			        scores = nx.pagerank(nx_graph)
			    except:
			        exit(1)			

			    ranked_sentences = sorted(((scores[i],s) for i,s in enumerate(sentences)), reverse=True)			

			    es = list()
			    for i in range(min(no_of_extracted_sentences, len(ranked_sentences))):
			        es.append(ranked_sentences[i][1])			

			    extracted_text = " ".join(es)			

			    return extracted_text

			def decode_sequence(input_seq):
			    # Encode the input as state vectors.
			    e_out, e_h, e_c = encoderModel.predict(input_seq)
			    
			    # Generate empty target sequence of length 1.
			    target_seq = np.zeros((1,1))
			    
			    # Populate the first word of target sequence with the start word.
			    target_seq[0, 0] = target_word_index['sostok']			

			    stop_condition = False
			    decoded_sentence = ''
			    while not stop_condition:
			      
			        output_tokens, h, c = decoderModel.predict([target_seq] + [e_out, e_h, e_c])			

			        # Sample a token
			        sampled_token_index = np.argmax(output_tokens[0, -1, :])
			        sampled_token = reverse_target_word_index[sampled_token_index]
			        
			        if(sampled_token!='eostok'):
			            decoded_sentence += ' '+sampled_token			

			        # Exit condition: either hit max length or find stop word.
			        if (sampled_token == 'eostok'  or len(decoded_sentence.split()) >= (max_summary_len-1)):
			            stop_condition = True			

			        # Update the target sequence (of length 1).
			        target_seq = np.zeros((1,1))
			        target_seq[0, 0] = sampled_token_index			

			        # Update internal states
			        e_h, e_c = h, c			

			    return decoded_sentence

			def seq2summary(input_seq):
			    newString=''
			    for i in input_seq:
			        if((i!=0 and i!=target_word_index['sostok']) and i!=target_word_index['eostok']):
			            newString=newString+reverse_target_word_index[i]+' '
			    return newString

			def seq2text(input_seq):
			    newString=''
			    for i in input_seq:
			        if(i!=0):
			            newString=newString+reverse_source_word_index[i]+' '
			    return newString


			extracted_text = extractText(abstract)
			cleaned_text = text_cleaner(extracted_text)			
			x_val_seq = x_tokenizer.texts_to_sequences([cleaned_text])
			text_encodings = pad_sequences(x_val_seq, maxlen=max_text_len, padding='post')[0]			
			resultTitle = decode_sequence(text_encodings.reshape(1,max_text_len))
			print(resultTitle)

			############	---------------		###############

			############	Code for Making the generated title Catchy 	###############

			def count(string):
				return sum(1 for c in string if c.isupper())			

			def catchy1(t,text,key):
				s = t.split()
				r = []
				flag = 0
				r.append(s[0])
				for i in range (1,len(s)):
					if s[i] == s[i-1]:
						continue
					else:
						r.append(s[i])
				for i in range (len(r)):
					r[i] = r[i][0].upper()+r[i][1:]
				stop_words = set(stopwords.words('english')) 
				text = re.sub(r'[^A-Za-z\s]',"",text)
				text = re.sub(r'https\S+', "", text)
				text = text.lstrip()
				text = text.strip()
				list1 = word_tokenize(text)
				fil_sent = [w for w in list1 if not w in stop_words]			

				final = []
				c = {}
				for i in fil_sent:
					if i[1:] != i[1:].lower() and count(i)>2:
						final.append(i)
						if i not in c:
							c[i]=1
						else:
							c[i] += 1
				ans = ""
				if len(c)!=0:
					max_value = max(c.values())
					last = [k for k,v in c.items() if v == max_value]
					length = []
					if len(last)>=1:
						for i in range (len(last)):
							if last[i] == last[i].upper():
								flag = 1
								ans = last[i]
								length.append(len(last[i]))
							else:
								length.append(len(last[i]))
						if ans == "":
							flag = 1
							ans = last[length.index(max(length))]
				final = []
				if ans == "":
					if len(key)>0:
						flag = 2
						pre = re.split('-|;|,|—|:',key)
						for i in range (len(pre)):
							if pre[i] != "" and pre[i] !="keywords" and pre[i]!="Keywords" and pre[i]!="KEYWORDS":
								final.append(pre[i])
						for i in range(len(final)):
							final[i] = final[i].lstrip()
							final[i] = final[i].rstrip()
						print(final)
						z = random.randrange(0,len(final))
						ans = final[z]
						ans = ans.replace(" ","")			

				ans1 = ""
				for i in range (len(r)):
					if i == 0:
						ans1 = ans1+r[i]
					else:
						ans1 = ans1+" "+r[i]
				if flag == 2:
					ans1 = "#"+ans+": "+ans1
				elif flag == 0:
					ans1 = ans1
				else:
					ans1 = ans+": "+ans1
				return (ans1)			

			catchy_title = catchy1(resultTitle,abstract,keywords)
			print("PREDICTED TITLE: ",catchy_title)

			# f3 = open("Predicted_Title.txt","w+")
			# f3.write("PREDICTED TITLE:- ",catchy_title)
			# f3.close()

			# 			#############	Code for Generating the metric for Catchiness	############

			# with open('xtokenizer.pickle', 'rb') as handle:
			#     x_tokenizer = pickle.load(handle)
			# with open('ytokenizer.pickle', 'rb') as handle:
			#     y_tokenizer = pickle.load(handle)						

			# dict1 = y_tokenizer.word_counts
			# dict2 = x_tokenizer.word_counts						

			# def merge_two_dicts(x, y):
			# 	z = x.copy()   
			# 	z.update(y)    
			# 	return z						

			# dict_new = merge_two_dicts(dict1,dict2)			

			# def catchyscore(ground,catchy):
			# 	stopwords = nltk.corpus.stopwords.words('english')				

			# 	ground_new = []
			# 	catchy_new = []				

			# 	ground = ground.split(" ")
			# 	catchy = catchy.split(" ")			

			# 	for i in ground:
			# 		if(i not in stopwords):
			# 			ground_new.append(i)			

			# 	for i in catchy:
			# 		if i not in stopwords:
			# 			catchy_new.append(i)				

			# 	ground_score = 0	# default values
			# 	catchy_score = random.choice([-1,0,1,2,3])	# default values			

			# 	for i in range(len(ground_new)):
			# 		if ground_new[i] in dict_new:
			# 			ground_score+=dict_new[i]				

			# 	for i in range(len(catchy_new)):
			# 		if catchy_new[i] in dict_new:
			# 			catchy_score+=dict_new[i]				

			# 	final = ground_score - catchy_score
			# 	return final		

			# if(abstract!=None):
			# 	catch_score = catchyscore(title,catchy_title)
			# 	print(catch_score)

			# #############	--------------	###############

			if(title):
				return render_template("mainpage.html", show_results=1, orgtit=orgtit , catchy_title=catchy_title, groundTruth=title)
			else:
				return render_template("mainpage.html", show_results=1, orgtit=orgtit , catchy_title=catchy_title)
		except:
			return render_template("mainpage.html", show_results=0)

	else:
		return render_template("mainpage.html", show_results=0)





@app.route("/index", methods=["GET","POST"])
def index():
 	return render_template("index2.html")



app.config['ENV'] = 'development'
app.config['DEBUG'] = True
app.config['TESTING'] = True
if __name__ == '__main__':
	app.run(debug=False)